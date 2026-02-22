"""
File Export Module for the Automated Book Generation System.
Supports DOCX, PDF, and TXT output formats.
"""

import os
from typing import List, Dict, Any, Optional
from datetime import datetime

import config


# =============================================================================
# BOOK EXPORTER
# =============================================================================

class BookExporter:
    """Exports book content to various file formats."""
    
    def __init__(self, output_dir: str = None):
        self.output_dir = output_dir or config.OUTPUT_DIRECTORY
        self._ensure_output_dir()
    
    def _ensure_output_dir(self):
        """Ensure output directory exists."""
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)
    
    def _sanitize_filename(self, name: str) -> str:
        """Sanitize filename for safe file system use."""
        # Remove or replace invalid characters
        invalid_chars = '<>:"/\\|?*'
        for char in invalid_chars:
            name = name.replace(char, '_')
        return name.strip()
    
    def _generate_filename(self, title: str, extension: str) -> str:
        """Generate unique filename."""
        sanitized_title = self._sanitize_filename(title)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        return f"{sanitized_title}_{timestamp}.{extension}"
    
    # =========================================================================
    # TEXT EXPORT
    # =========================================================================
    
    def export_to_txt(
        self,
        title: str,
        chapters: List[Dict[str, Any]],
        outline: str = None
    ) -> str:
        """Export book to plain text file."""
        filename = self._generate_filename(title, "txt")
        filepath = os.path.join(self.output_dir, filename)
        
        with open(filepath, 'w', encoding='utf-8') as f:
            # Title
            f.write("=" * 60 + "\n")
            f.write(f"{title.upper()}\n")
            f.write("=" * 60 + "\n\n")
            
            # Optional outline
            if outline:
                f.write("TABLE OF CONTENTS\n")
                f.write("-" * 40 + "\n")
                f.write(outline + "\n\n")
                f.write("=" * 60 + "\n\n")
            
            # Chapters
            for chapter in chapters:
                f.write(f"\nCHAPTER {chapter['chapter_number']}: {chapter['title']}\n")
                f.write("-" * 40 + "\n\n")
                f.write(chapter.get('content', '') + "\n\n")
                f.write("=" * 60 + "\n")
            
            # Footer
            f.write(f"\n\nGenerated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        
        print(f"✓ Exported to TXT: {filepath}")
        return filepath
    
    # =========================================================================
    # DOCX EXPORT
    # =========================================================================
    
    def export_to_docx(
        self,
        title: str,
        chapters: List[Dict[str, Any]],
        outline: str = None
    ) -> str:
        """Export book to Word document (.docx)."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.style import WD_STYLE_TYPE
        except ImportError:
            raise ImportError("Please install python-docx: pip install python-docx")
        
        filename = self._generate_filename(title, "docx")
        filepath = os.path.join(self.output_dir, filename)
        
        doc = Document()
        
        # Title page
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(28)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add some space
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Generated date
        date_para = doc.add_paragraph()
        date_para.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()
        
        # Table of Contents (if outline provided)
        if outline:
            toc_heading = doc.add_heading("Table of Contents", level=1)
            doc.add_paragraph(outline)
            doc.add_page_break()
        
        # Chapters
        for chapter in chapters:
            # Chapter heading
            chapter_heading = doc.add_heading(
                f"Chapter {chapter['chapter_number']}: {chapter['title']}",
                level=1
            )
            
            # Chapter content
            content = chapter.get('content', '')
            
            # Split content into paragraphs
            paragraphs = content.split('\n\n')
            for para_text in paragraphs:
                if para_text.strip():
                    # Check if it's a subheading (starts with ## or similar)
                    if para_text.strip().startswith('##'):
                        heading_text = para_text.strip().lstrip('#').strip()
                        doc.add_heading(heading_text, level=2)
                    elif para_text.strip().startswith('#'):
                        heading_text = para_text.strip().lstrip('#').strip()
                        doc.add_heading(heading_text, level=2)
                    else:
                        doc.add_paragraph(para_text.strip())
            
            doc.add_page_break()
        
        # Save document
        doc.save(filepath)
        print(f"✓ Exported to DOCX: {filepath}")
        return filepath
    
    # =========================================================================
    # PDF EXPORT
    # =========================================================================
    
    def export_to_pdf(
        self,
        title: str,
        chapters: List[Dict[str, Any]],
        outline: str = None
    ) -> str:
        """Export book to PDF file."""
        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
            from reportlab.platypus import Table, TableStyle
        except ImportError:
            raise ImportError("Please install reportlab: pip install reportlab")
        
        filename = self._generate_filename(title, "pdf")
        filepath = os.path.join(self.output_dir, filename)
        
        # Create document
        doc = SimpleDocTemplate(
            filepath,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )
        
        # Styles
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'BookTitle',
            parent=styles['Heading1'],
            fontSize=28,
            spaceAfter=30,
            alignment=1  # Center
        )
        
        chapter_title_style = ParagraphStyle(
            'ChapterTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceBefore=20,
            spaceAfter=20
        )
        
        body_style = ParagraphStyle(
            'BookBody',
            parent=styles['Normal'],
            fontSize=11,
            leading=16,
            spaceBefore=6,
            spaceAfter=6
        )
        
        # Build content
        story = []
        
        # Title page
        story.append(Spacer(1, 2*inch))
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, inch))
        story.append(Paragraph(
            f"Generated: {datetime.now().strftime('%B %d, %Y')}",
            ParagraphStyle('Date', parent=styles['Normal'], alignment=1)
        ))
        story.append(PageBreak())
        
        # Table of Contents
        if outline:
            story.append(Paragraph("Table of Contents", chapter_title_style))
            story.append(Spacer(1, 0.25*inch))
            for line in outline.split('\n'):
                if line.strip():
                    story.append(Paragraph(line, body_style))
            story.append(PageBreak())
        
        # Chapters
        for chapter in chapters:
            # Chapter title
            story.append(Paragraph(
                f"Chapter {chapter['chapter_number']}: {chapter['title']}",
                chapter_title_style
            ))
            story.append(Spacer(1, 0.25*inch))
            
            # Chapter content
            content = chapter.get('content', '')
            paragraphs = content.split('\n\n')
            
            for para_text in paragraphs:
                if para_text.strip():
                    # Handle markdown-style headers
                    if para_text.strip().startswith('##'):
                        heading_text = para_text.strip().lstrip('#').strip()
                        story.append(Paragraph(heading_text, styles['Heading2']))
                    elif para_text.strip().startswith('#'):
                        heading_text = para_text.strip().lstrip('#').strip()
                        story.append(Paragraph(heading_text, styles['Heading2']))
                    else:
                        # Regular paragraph - escape special characters
                        safe_text = para_text.strip()
                        safe_text = safe_text.replace('&', '&amp;')
                        safe_text = safe_text.replace('<', '&lt;')
                        safe_text = safe_text.replace('>', '&gt;')
                        try:
                            story.append(Paragraph(safe_text, body_style))
                        except:
                            # If paragraph fails, add as plain text
                            story.append(Paragraph(
                                safe_text[:500] + "..." if len(safe_text) > 500 else safe_text,
                                body_style
                            ))
            
            story.append(PageBreak())
        
        # Build PDF
        doc.build(story)
        print(f"✓ Exported to PDF: {filepath}")
        return filepath
    
    # =========================================================================
    # EXPORT ALL FORMATS
    # =========================================================================
    
    def export_all(
        self,
        title: str,
        chapters: List[Dict[str, Any]],
        outline: str = None,
        formats: List[str] = None
    ) -> Dict[str, str]:
        """Export book to all specified formats."""
        formats = formats or config.OUTPUT_FORMATS
        results = {}
        
        for fmt in formats:
            try:
                if fmt.lower() == 'txt':
                    results['txt'] = self.export_to_txt(title, chapters, outline)
                elif fmt.lower() == 'docx':
                    results['docx'] = self.export_to_docx(title, chapters, outline)
                elif fmt.lower() == 'pdf':
                    results['pdf'] = self.export_to_pdf(title, chapters, outline)
                else:
                    print(f"⚠ Unsupported format: {fmt}")
            except Exception as e:
                print(f"✗ Error exporting to {fmt}: {e}")
                results[fmt] = None
        
        return results
    
    # =========================================================================
    # CHAPTER EXPORT (Individual)
    # =========================================================================
    
    def export_chapter_to_txt(
        self,
        book_title: str,
        chapter_number: int,
        chapter_title: str,
        content: str
    ) -> str:
        """Export individual chapter to text file."""
        filename = f"{self._sanitize_filename(book_title)}_chapter_{chapter_number}.txt"
        filepath = os.path.join(self.output_dir, filename)
        
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(f"CHAPTER {chapter_number}: {chapter_title}\n")
            f.write("=" * 40 + "\n\n")
            f.write(content)
            f.write(f"\n\n---\nExported: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        return filepath
    
    # =========================================================================
    # OUTLINE EXPORT
    # =========================================================================
    
    def export_outline_to_txt(self, book_title: str, outline: str) -> str:
        """Export outline to text file."""
        filename = f"{self._sanitize_filename(book_title)}_outline.txt"
        filepath = os.path.join(self.output_dir, filename)
        
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(f"OUTLINE: {book_title}\n")
            f.write("=" * 40 + "\n\n")
            f.write(outline)
            f.write(f"\n\n---\nExported: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        return filepath


# =============================================================================
# CONVENIENCE FUNCTION
# =============================================================================

def get_exporter() -> BookExporter:
    """Get book exporter instance."""
    return BookExporter()
